import sys
import re
import docx.document
import random



def makeTemplate():
    doc: docx.document.Document = docx.Document("口算与笔算.docx")
    patten = re.compile(f"\d")
    for p in doc.paragraphs:
        print(p.text)
        for run in p.runs:
            run.text = patten.sub("{}", run.text)
        print(p.text)
    doc.save("口算与笔算_模板.docx")

def createNewDoc(fileName:str):
    doc: docx.document.Document = docx.Document("口算与笔算_模板.docx")
    patten = re.compile("\{\}")
    for p in doc.paragraphs:
        for run in p.runs:
            mathes = patten.findall(run.text)
            for i in range(len(mathes)):
                mathes[i] = random.randrange(1, 9)
            run.text = run.text.format(*mathes)
        print(p.text)
    doc.save(fileName)

# makeTemplate()
for i in range(11, 25):
    createNewDoc(f"口算与笔算（{i}）.docx")
sys.exit()

